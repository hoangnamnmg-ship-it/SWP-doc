
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_docx_from_md(md_path, output_path):
    document = Document()
    
    # Simple markdown parser
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    table_mode = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Heading 1
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        # Heading 2
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        # Heading 3
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        # Horizontal Rule
        elif line.startswith("---"):
            document.add_paragraph("_" * 50)
        # Bullet list
        elif line.startswith("* "):
            document.add_paragraph(line[2:], style='List Bullet')
        # Blockquote/Note
        elif line.startswith("> "):
             p = document.add_paragraph()
             runner = p.add_run(line[2:])
             runner.italic = True
        # Table handling
        elif line.startswith("|"):
            if not table_mode:
                table_mode = True
                table_data = []
            
            # Remove leading/trailing pipes and split
            row_content = [cell.strip() for cell in line.strip("|").split("|")]
            
            # Skip separator lines like |---|---|
            if "---" in row_content[0]:
                continue
                
            table_data.append(row_content)
        
        # End of table or normal text
        else:
            if table_mode:
                if table_data:
                    # Create table
                    try:
                        table = document.add_table(rows=1, cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(table_data[0]):
                            hdr_cells[i].text = header
                            # Bold headers
                            for run in hdr_cells[i].paragraphs[0].runs:
                                run.bold = True
                        
                        for row in table_data[1:]:
                            row_cells = table.add_row().cells
                            for i, cell in enumerate(row):
                                if i < len(row_cells):
                                    row_cells[i].text = cell
                    except Exception as e:
                        print(f"Error creating table: {e}")
                    
                    document.add_paragraph() # Spacer
                table_mode = False
                table_data = []

            if line:
                # Handle bold in normal text (**text**) - simplified
                p = document.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    document.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    # Convert Vietnamese version
    md_file_vn = r"E:\SWP doc\vision and scope vn.md"
    docx_file_vn = r"E:\SWP doc\vision and scope vn.docx"
    create_docx_from_md(md_file_vn, docx_file_vn)
